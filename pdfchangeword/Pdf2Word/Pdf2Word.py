#coding=utf-8

'''
Pdf 转 word
'''

import os
import time
import sys
from configparser import ConfigParser
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor

from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document


# pdf_folder = r'./Pdfdocument'
# word_folder = r'./Worddocument'

def read_from_pdf(file_path):  # 读取pdf文件
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()  # 创建对象
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content  # 返回pdf中的内容


def remove_control_characters(content):  # 移除内容中的32以下的字符
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def save_text_to_word(content, file_path):  # 将文本流保存为word文档
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def pdf_to_word(pdf_file_path, word_file_path):  # 读取内容与写入内容
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)


def main():
    config_parser = ConfigParser()
    config_parser.read('config.cfg')
    config = config_parser['default']
    tasks = []
    dir_lists = [d for d in os.listdir('./')]

    if config['pdf_folder'] not in dir_lists or config['word_folder'] not in dir_lists:
        os.mkdir(config['pdf_folder'])
        os.mkdir(config['word_folder'])
        print("创建文件夹成功，请添加pdf文件!")
        time.sleep(2)
        sys.exit()

    # with ProcessPoolExecutor(max_workers=int(config['max_worker'])) as executor:  # 调用线程执行读取pdf与写入生成word
    try:
        for file in os.listdir('./' + config['pdf_folder'] + '/'):
            extension_name = os.path.splitext(file)[1]  # 判断文件是否是pdf格式
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(file)[0]
            pdf_file = './' + config['pdf_folder'] + '/' + file
            word_file = './' + config['word_folder'] + '/' + file_name + '.docx'
            print('正在处理: ', file)
            time.sleep(5)
            pdf_to_word(pdf_file, word_file)
            # result = executor.submit(pdf_to_word, pdf_file, word_file)
            # result = executor.map(pdf_to_word, pdf_file, word_file)
            # tasks.append(result)
    except Exception as ret:
        print(ret)
    else:
        print("--处理完毕--")
        time.sleep(5)
        sys.exit(0)

    # while True:
    #     exit_flag = True
    #     for task in tasks:
    #         if not task.done():
    #             exit_flag = False
    #     if exit_flag:
    #         print('----关闭应用----')
    #         break
    #         sys.exit()  # 正常退出

    # for task in tasks:
    #     if not task.done():
    #         continue
    #     else:
    #         print('----关闭应用----')
    #         time.sleep(5)
    #         break
    #         exit(0)  # 正常退出


if __name__ == '__main__':
    try:
        main()
    except Exception as ex:
        print(ex)
        input("请检查异常")
